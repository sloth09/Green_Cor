#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate paper.docx from paper_final.md using python-docx (no COM automation)."""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ---- Figure configuration ----
FIGURES_DIR = r"D:\code\Green_Cor\results\paper_figures"

# Figure label -> (filename, caption)
FIGURE_MAP = {
    "Fig. 1": ("D7_cycle_time.png", "Cycle time components for three supply chain configurations"),
    "Fig. 2": ("D1_npc_vs_shuttle.png", "Net present cost vs shuttle size for all cases"),
    "Fig. 3": ("D10_case_npc_comparison.png", "NPC comparison across cases at optimal configurations"),
    "Fig. 4": ("D6_cost_breakdown.png", "Cost component breakdown (CAPEX/OPEX)"),
    "Fig. 5": ("D9_lco_comparison.png", "Levelized cost of ammonia bunkering by case"),
    "Fig. 6": ("D2_yearly_cost_evolution.png", "Annual cost evolution (2030-2050)"),
    "Fig. 7": ("D8_fleet_evolution.png", "Fleet size evolution over planning horizon"),
    "Fig. 8": ("D3_yearly_fleet_demand.png", "Annual bunkering demand and fleet response"),
    "Fig. 9": ("D5_yearly_utilization.png", "Fleet utilization rates over time"),
    "Fig. 10": ("Fig7_tornado_deterministic.png", "Tornado diagram: parametric sensitivity of NPC"),
    "Fig. 11": ("Fig8_fuel_price_sensitivity.png", "Fuel price sensitivity: NPC and LCO response"),
    "Fig. 12": ("Fig9_breakeven_distance.png", "Break-even distance analysis: Case 1 vs Case 2"),
    "Fig. 13": ("Fig10_demand_scenarios.png", "Demand scenario analysis: NPC and LCO"),
    "Fig. 14": ("S7_pump_sensitivity.png", "Pump rate sensitivity analysis"),
    "Fig. 15": ("Fig11_discount_rate_sensitivity.png", "Discount rate sensitivity: NPC and LCOA across three cases"),
    "Fig. 16": ("Fig12_discount_rate_fleet.png", "Fleet evolution under different discount rates"),
    "Fig. 17": ("Fig13_yang_lam_service_time.png", "Service time comparison: MILP model vs Yang and Lam DES model"),
    "Fig. 18": ("Fig14_yang_lam_sensitivity.png", "Flow rate sensitivity comparison: MILP vs DES model"),
    "Fig. S1": ("D12_npc_heatmaps.png", "NPC sensitivity heatmap (shuttle size x pump rate)"),
    "Fig. S2": ("D11_top_configurations.png", "Top configurations ranked by NPC"),
    "Fig. S3": ("D4_yearly_cycles.png", "Annual cycle count evolution"),
    "Fig. S4": ("FigS4_twoway_deterministic.png", "Two-way sensitivity: fuel price x bunker volume"),
    "Fig. S5": ("FigS5_bunker_volume_sensitivity.png", "Bunker volume sensitivity: NPC and LCO response"),
}


def parse_markdown(md_path):
    """Parse paper_final.md into structured sections."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    sections = []
    current_section = None
    buffer = []

    for line in lines:
        line = line.rstrip("\n")

        # Skip horizontal rules
        if line.strip() == "---":
            if buffer and current_section is not None:
                sections.append((current_section, "\n".join(buffer)))
                buffer = []
            continue

        # Heading detection
        h1 = re.match(r"^# (.+)$", line)
        h2 = re.match(r"^## (.+)$", line)
        h3 = re.match(r"^### (.+)$", line)

        if h1:
            if buffer and current_section is not None:
                sections.append((current_section, "\n".join(buffer)))
                buffer = []
            current_section = ("h1", h1.group(1))
        elif h2:
            if buffer and current_section is not None:
                sections.append((current_section, "\n".join(buffer)))
                buffer = []
            current_section = ("h2", h2.group(1))
        elif h3:
            if buffer and current_section is not None:
                sections.append((current_section, "\n".join(buffer)))
                buffer = []
            current_section = ("h3", h3.group(1))
        else:
            buffer.append(line)

    if buffer and current_section is not None:
        sections.append((current_section, "\n".join(buffer)))

    return sections


def add_rich_paragraph(doc, text, style=None, bold=False, font_size=None, space_after=None):
    """Add a paragraph with inline bold/italic formatting."""
    p = doc.add_paragraph(style=style)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)

    # Split text by bold markers **...**
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            run = p.add_run(part)

        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)

    return p


def insert_figure(doc, fig_label, caption, img_path):
    """Insert figure image with centered caption."""
    if not os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{fig_label}: Image not found at {os.path.basename(img_path)}]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(180, 0, 0)
        return

    doc.add_picture(img_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(f"{fig_label}. {caption}")
    run.font.size = Pt(9)
    run.italic = True
    cap_p.paragraph_format.space_after = Pt(12)


def _insert_figures_from_line(doc, line, inserted_figures):
    """Check line for figure references and insert images after first mention."""
    refs = re.findall(r'Fig\. (S?\d+)', line)
    for ref_num in refs:
        fig_key = f"Fig. {ref_num}"
        if fig_key in FIGURE_MAP and fig_key not in inserted_figures:
            filename, caption = FIGURE_MAP[fig_key]
            img_path = os.path.join(FIGURES_DIR, filename)
            insert_figure(doc, fig_key, caption, img_path)
            inserted_figures.add(fig_key)


def add_table_from_md(doc, table_text):
    """Parse markdown table and add to document."""
    lines = [l.strip() for l in table_text.strip().split("\n") if l.strip()]
    # Filter out separator lines (|---|---|)
    data_lines = []
    for line in lines:
        stripped = line.strip("|").strip()
        if re.match(r"^[\s\-:|]+$", stripped):
            continue
        data_lines.append(line)

    if len(data_lines) < 2:
        return None

    # Parse cells
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                    if i == 0:
                        for run in paragraph.runs:
                            run.bold = True

    return table


def process_body(doc, body_text, inserted_figures):
    """Process body text, handling tables, figures, bold text, and regular paragraphs."""
    lines = body_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Detect table start
        if "|" in line and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            add_table_from_md(doc, "\n".join(table_lines))
            doc.add_paragraph()  # spacing after table
            continue

        # Bold-only line (like **Table 1: ...** or **Gap 1:**)
        if line.strip().startswith("**") and line.strip().count("**") >= 2:
            add_rich_paragraph(doc, line.strip(), font_size=11, space_after=4)
            _insert_figures_from_line(doc, line, inserted_figures)
            i += 1
            continue

        # Regular paragraph
        add_rich_paragraph(doc, line.strip(), font_size=11, space_after=6)
        _insert_figures_from_line(doc, line, inserted_figures)
        i += 1


def main():
    md_path = r"D:\code\Green_Cor\docs\paper\v4\paper_final.md"
    out_path = r"D:\code\Green_Cor\docs\paper\v4\paper.docx"

    sections = parse_markdown(md_path)

    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # -- Process sections --
    inserted_figures = set()

    for heading, body in sections:
        level, title = heading

        if level == "h1":
            # Paper title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_after = Pt(12)

        elif level == "h2":
            p = doc.add_heading(title, level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif level == "h3":
            p = doc.add_heading(title, level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Process body content
        if body.strip():
            process_body(doc, body, inserted_figures)

    # -- Insert supplementary figures not referenced in body text --
    remaining = [k for k in FIGURE_MAP if k not in inserted_figures]
    if remaining:
        p = doc.add_heading("Supplementary Figures", level=1)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        for fig_key in remaining:
            filename, caption = FIGURE_MAP[fig_key]
            img_path = os.path.join(FIGURES_DIR, filename)
            insert_figure(doc, fig_key, caption, img_path)
            inserted_figures.add(fig_key)

    doc.save(out_path)
    print(f"[OK] {out_path} saved ({len(inserted_figures)} figures inserted)")

    # -- Verification: count sections and tables --
    verify_docx(out_path, md_path)


def verify_docx(docx_path, md_path):
    """Verify DOCX content against markdown source."""
    doc = Document(docx_path)

    # Count paragraphs and tables
    n_paragraphs = len(doc.paragraphs)
    n_tables = len(doc.tables)

    # Extract all text from DOCX
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_text += " " + cell.text

    # Read markdown
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Key content checks - updated for v4 paper (pump 500 m3/h)
    checks = [
        # Structure
        ("Title", "Optimal Ammonia Bunkering Infrastructure"),
        ("Section 1", "1. Introduction"),
        ("Subsection 1.1", "Related Work"),
        ("Subsection 1.2", "Research Gaps"),
        ("Section 2", "2. Methodology"),
        ("Section 3", "3. Results and Analysis"),
        ("Section 4", "4. Discussion"),
        ("Section 5", "5. Conclusions"),
        # Abstract
        ("Abstract keyword", "ammonia bunkering"),
        # Core numbers (v4: pump 500 m3/h)
        ("Case 1 NPC", "447.53"),
        ("Case 2 NPC", "906.80"),
        ("Case 3 NPC", "1,094.12"),
        ("LCOA Case 1", "1.90"),
        ("Demand robustness", "3.7%"),
        ("CAPEX scaling swing", "391.82"),
        ("Annuity factor", "10.8355"),
        # Methodology depth
        ("Notation table", "Decision variables"),
        ("Eq. numbering", "(15)"),
        ("CAPEX formula", "61.5"),
        ("Worked example", "3.87"),
        ("SFOC table", "4-stroke high-speed"),
        ("Pump rate 500", "500"),
        ("Assumption A1", "A1"),
        ("Assumption A6", "A6"),
        # Results depth
        ("Tornado Table 7", "CAPEX Scaling"),
        ("Demand Table 8", "VeryHigh"),
        ("Cost breakdown table", "Bunkering CAPEX"),
        ("Utilization sawtooth", "sawtooth"),
        # Discussion depth
        ("Recommendation 1", "Recommendation 1"),
        ("Recommendation 3", "Recommendation 3"),
        ("Limitation L1", "L1"),
        ("Limitation L6", "L6"),
        ("Future work F1", "F1"),
        ("Future work F5", "F5"),
        ("Yang and Lam comparison", "51.3%"),
        # Literature
        ("Lit - Yang Lam", "Yang and Lam"),
        # References
        ("Ref first", "Al-Enazi"),
        # Figures
        ("Figure list", "Fig. 1"),
        ("Supplementary", "Fig. S5"),
        # Draft artifact check
        ("No draft artifact", "wait,"),
    ]

    print(f"\n{'='*60}")
    print(f"DOCX Verification Report")
    print(f"{'='*60}")
    print(f"  Paragraphs: {n_paragraphs}")
    print(f"  Tables:     {n_tables}")
    print(f"{'='*60}")

    # Count words
    word_count = len(docx_text.split())
    print(f"  Words:      ~{word_count}")
    print(f"{'='*60}")

    passed = 0
    failed = 0
    for label, keyword in checks:
        if label == "No draft artifact":
            # This should NOT be found
            if keyword not in docx_text:
                print(f"  [PASS] {label}: '{keyword}' correctly absent")
                passed += 1
            else:
                print(f"  [FAIL] {label}: '{keyword}' STILL PRESENT (draft artifact)")
                failed += 1
        else:
            if keyword in docx_text:
                print(f"  [PASS] {label}")
                passed += 1
            else:
                print(f"  [FAIL] {label}: '{keyword}' NOT found")
                failed += 1

    print(f"{'='*60}")
    print(f"  Result: {passed}/{passed+failed} checks passed")
    if failed == 0:
        print(f"  [OK] DOCX content matches markdown source")
    else:
        print(f"  [WARN] {failed} checks failed - review needed")


if __name__ == "__main__":
    main()
