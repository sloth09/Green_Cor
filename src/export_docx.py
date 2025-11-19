"""
Export optimization results to Word document format.
Creates a professional report with analysis summary and key findings.
"""

from pathlib import Path
from typing import Dict
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordExporter:
    """Export MILP results to Word document."""

    def __init__(self, config: Dict):
        """
        Initialize Word exporter.

        Args:
            config: Configuration dictionary
        """
        self.config = config
        self.case_name = config.get("case_name", "Unknown")
        self.case_id = config.get("case_id", "unknown")

    def export_report(
        self,
        scenario_df: pd.DataFrame,
        yearly_df: pd.DataFrame,
        output_path: Path = None
    ) -> Path:
        """
        Export results to Word document.

        Args:
            scenario_df: Scenario summary DataFrame
            yearly_df: Yearly results DataFrame
            output_path: Output directory (default: results/)

        Returns:
            Path to created Word file
        """
        if output_path is None:
            output_path = Path("results")

        output_path.mkdir(parents=True, exist_ok=True)

        filename = f"MILP_Report_{self.case_id}.docx"
        filepath = output_path / filename

        # Create document
        doc = Document()

        # Add content
        self._add_title_page(doc)
        self._add_executive_summary(doc, scenario_df, yearly_df)
        self._add_case_description(doc)
        self._add_optimal_solution(doc, scenario_df)
        self._add_time_analysis(doc, scenario_df)
        self._add_scenario_analysis(doc, scenario_df)
        self._add_cost_breakdown(doc, scenario_df)
        self._add_appendix(doc, scenario_df, yearly_df)

        # Save
        doc.save(filepath)
        print(f"Word document export completed: {filepath}")

        return filepath

    def _add_title_page(self, doc: Document) -> None:
        """Add title page."""
        # Title
        title = doc.add_heading("Ammonia Bunkering Infrastructure Optimization", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Case name
        subtitle = doc.add_heading(self.case_name, level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Metadata
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = "Light Grid Accent 1"

        meta_data = [
            ("Case ID", self.case_id),
            ("Report Date", datetime.now().strftime("%Y-%m-%d")),
            ("Time Period", f"{self.config['time_period']['start_year']}-{self.config['time_period']['end_year']}"),
            ("Analysis Tool", "Green Corridor MILP Optimizer v2.0"),
        ]

        for idx, (key, value) in enumerate(meta_data):
            meta_table.rows[idx].cells[0].text = key
            meta_table.rows[idx].cells[1].text = str(value)

    def _add_executive_summary(
        self,
        doc: Document,
        scenario_df: pd.DataFrame,
        yearly_df: pd.DataFrame
    ) -> None:
        """Add executive summary section."""
        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)

        # Key findings
        best_scenario = scenario_df.nsmallest(1, "NPC_Total_USDm").iloc[0]

        summary_text = f"""
This report presents the optimization results for {self.case_name}.
The analysis evaluates {len(scenario_df)} different configurations of shuttle vessels and pump systems
over a {self.config['time_period']['end_year'] - self.config['time_period']['start_year']}-year period
(2030-2050) to minimize the Net Present Cost (NPC).
"""
        doc.add_paragraph(summary_text)

        # Key metrics table
        doc.add_heading("Key Metrics", level=2)
        metrics_table = doc.add_table(rows=6, cols=2)
        metrics_table.style = "Light Grid Accent 1"

        metrics = [
            ("Total Feasible Scenarios", len(scenario_df)),
            ("Optimal Shuttle Size (m³)", int(best_scenario["Shuttle_Size_cbm"])),
            ("Optimal Pump Flow (m³/h)", int(best_scenario["Pump_Size_m3ph"])),
            ("Minimum NPC (M USD)", f"{best_scenario['NPC_Total_USDm']:.2f}"),
            ("Annual Demand (2050, m³)", f"{yearly_df[yearly_df['Year']==2050]['Demand_m3'].iloc[0]:,.0f}"),
        ]

        for idx, (metric, value) in enumerate(metrics):
            metrics_table.rows[idx].cells[0].text = metric
            metrics_table.rows[idx].cells[1].text = str(value)

    def _add_case_description(self, doc: Document) -> None:
        """Add case description section."""
        doc.add_page_break()
        doc.add_heading("Case Description", level=1)

        # Route information
        route = self.config.get("route", {})
        if route:
            doc.add_heading("Route Information", level=2)
            doc.add_paragraph(f"Source: {route.get('source', 'N/A')}")
            doc.add_paragraph(f"Destination: {route.get('destination', 'N/A')}")
            doc.add_paragraph(
                f"Distance: {route.get('distance_nautical_miles', 'N/A')} nautical miles"
            )
            doc.add_paragraph(f"Travel Time: {self.config['operations']['travel_time_hours']:.2f} hours")

        # Operational parameters
        doc.add_heading("Operational Parameters", level=2)
        ops_table = doc.add_table(rows=6, cols=2)
        ops_table.style = "Light Grid Accent 1"

        ops = [
            ("Shuttle Sizes Available (m³)", ", ".join(map(str, self.config['shuttle']['available_sizes_cbm'][:5])) + "..."),
            ("Pump Flow Rates (m³/h)", ", ".join(map(str, self.config['pumps']['available_flow_rates'][:5])) + "..."),
            ("Tank Storage", "Enabled" if self.config['tank_storage']['enabled'] else "Disabled"),
            ("Tank Size (tons)", self.config['tank_storage']['size_tons']),
            ("Setup Time (hours)", self.config['operations']['setup_time_hours']),
            ("Max Annual Hours", self.config['operations']['max_annual_hours_per_vessel']),
        ]

        for idx, (param, value) in enumerate(ops):
            ops_table.rows[idx].cells[0].text = param
            ops_table.rows[idx].cells[1].text = str(value)

    def _add_optimal_solution(self, doc: Document, scenario_df: pd.DataFrame) -> None:
        """Add optimal solution section."""
        doc.add_page_break()
        doc.add_heading("Optimal Solution", level=1)

        best = scenario_df.nsmallest(1, "NPC_Total_USDm").iloc[0]

        doc.add_paragraph(
            f"The optimization identifies Shuttle Size {int(best['Shuttle_Size_cbm'])} m³ "
            f"with Pump Flow {int(best['Pump_Size_m3ph'])} m³/h as the optimal configuration."
        )

        # Cost breakdown
        doc.add_heading("Cost Breakdown", level=2)
        cost_table = doc.add_table(rows=10, cols=2)
        cost_table.style = "Light Grid Accent 1"

        costs = [
            ("Total NPC (M USD)", f"{best['NPC_Total_USDm']:.2f}"),
            ("Shuttle CAPEX (M USD)", f"{best['NPC_Shuttle_CAPEX_USDm']:.2f}"),
            ("Bunkering CAPEX (M USD)", f"{best['NPC_Bunkering_CAPEX_USDm']:.2f}"),
            ("Terminal CAPEX (M USD)", f"{best['NPC_Terminal_CAPEX_USDm']:.2f}"),
            ("Shuttle fOPEX (M USD)", f"{best['NPC_Shuttle_fOPEX_USDm']:.2f}"),
            ("Bunkering fOPEX (M USD)", f"{best['NPC_Bunkering_fOPEX_USDm']:.2f}"),
            ("Shuttle vOPEX (M USD)", f"{best['NPC_Shuttle_vOPEX_USDm']:.2f}"),
            ("Bunkering vOPEX (M USD)", f"{best['NPC_Bunkering_vOPEX_USDm']:.2f}"),
            ("Terminal vOPEX (M USD)", f"{best['NPC_Terminal_vOPEX_USDm']:.2f}"),
        ]

        for idx, (cost_item, value) in enumerate(costs):
            cost_table.rows[idx].cells[0].text = cost_item
            cost_table.rows[idx].cells[1].text = str(value)

    def _add_time_analysis(self, doc: Document, scenario_df: pd.DataFrame) -> None:
        """Add operation time analysis section."""
        doc.add_page_break()
        doc.add_heading("운항 시간 분석 (Time Structure Analysis)", level=1)

        best = scenario_df.nsmallest(1, "NPC_Total_USDm").iloc[0]

        doc.add_heading("최적 시나리오의 운항 시간 구성", level=2)

        intro = (
            f"최적 시나리오 (Shuttle {int(best['Shuttle_Size_cbm'])} m³, "
            f"Pump {int(best['Pump_Size_m3ph'])} m³/h)의 1회 왕복 운항 시간을 상세히 분석합니다. "
            f"총 사이클 시간은 {best['Cycle_Duration_hr']:.2f}시간이며, "
            f"이는 연간 최대 {best['Annual_Cycles_Max']:.0f}회의 운항을 가능하게 합니다."
        )
        doc.add_paragraph(intro)

        # Time breakdown table
        doc.add_heading("시간 구성 요소 분해", level=3)

        time_table = doc.add_table(rows=8, cols=3)
        time_table.style = "Light Grid Accent 1"

        # Header
        header_cells = time_table.rows[0].cells
        header_cells[0].text = "시간 구성 요소"
        header_cells[1].text = "시간 (h)"
        header_cells[2].text = "비율 (%)"

        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Time components
        basic_cycle = best.get("Basic_Cycle_Duration_hr", 0)
        components = [
            ("육상 연료 적재", best.get("Shore_Loading_hr", 0)),
            ("편도 항해", best.get("Travel_Outbound_hr", 0)),
            ("호스 연결 및 퍼징", best.get("Setup_Inbound_hr", 0)),
            ("벙커링 펌핑", best.get("Pumping_Per_Vessel_hr", 0)),
            ("호스 분리 및 퍼징", best.get("Setup_Outbound_hr", 0)),
            ("복귀 항해", best.get("Travel_Return_hr", 0)),
        ]

        for idx, (name, hours) in enumerate(components, 1):
            row_cells = time_table.rows[idx].cells
            row_cells[0].text = name
            row_cells[1].text = f"{hours:.2f}"
            percentage = (hours / basic_cycle * 100) if basic_cycle > 0 else 0
            row_cells[2].text = f"{percentage:.1f}%"

        # Total row
        total_row_cells = time_table.rows[7].cells
        total_row_cells[0].text = "【총 사이클 시간】"
        total_row_cells[1].text = f"{best['Cycle_Duration_hr']:.2f}"
        total_row_cells[2].text = "100%"

        for cell in total_row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Operating metrics
        doc.add_heading("연간 운영 지표", level=3)

        metrics_table = doc.add_table(rows=5, cols=2)
        metrics_table.style = "Light Grid Accent 1"

        metrics = [
            ("연간 최대 항차", f"{best['Annual_Cycles_Max']:.0f} 회"),
            ("연간 공급 용량", f"{best['Annual_Supply_m3']:.0f} m³"),
            ("시간 활용도", f"{best['Time_Utilization_Ratio_percent']:.1f}%"),
            ("셔틀당 평균 일정", f"{365/best['Annual_Cycles_Max']:.1f} 일" if best['Annual_Cycles_Max'] > 0 else "N/A"),
        ]

        for idx, (metric_name, value) in enumerate(metrics, 0):
            metrics_table.rows[idx].cells[0].text = metric_name
            metrics_table.rows[idx].cells[1].text = str(value)

        # Analysis
        doc.add_heading("시간 분석", level=3)

        # Find which component takes the most time
        max_component = max(components, key=lambda x: x[1])
        max_percentage = (max_component[1] / basic_cycle * 100) if basic_cycle > 0 else 0

        analysis_text = (
            f"벙커링 펌핑이 전체 사이클의 {max_percentage:.1f}%를 차지하여, "
            f"펌프 크기 선택이 가장 중요한 최적화 포인트입니다. "
            f"기본 사이클 시간(육상 제외)은 {best['Basic_Cycle_Duration_hr']:.2f}시간이며, "
            f"육상 연료 적재에 추가로 {best.get('Shore_Loading_hr', 0):.2f}시간이 소요됩니다. "
            f"이 시간 구조는 연간 {best['Annual_Cycles_Max']:.0f}회의 운항을 가능하게 하며, "
            f"총 20년간 약 {best['Annual_Supply_m3'] * 20 / 1e6:.1f}백만 m³의 암모니아 연료를 공급할 수 있습니다."
        )
        doc.add_paragraph(analysis_text)

    def _add_scenario_analysis(self, doc: Document, scenario_df: pd.DataFrame) -> None:
        """Add scenario analysis section."""
        doc.add_page_break()
        doc.add_heading("Scenario Analysis", level=1)

        doc.add_paragraph(
            "This section presents the top scenarios ranked by Net Present Cost (NPC)."
        )

        # Top 10 scenarios
        doc.add_heading("Top 10 Scenarios", level=2)
        top10 = scenario_df.nsmallest(10, "NPC_Total_USDm")

        table = doc.add_table(rows=len(top10) + 1, cols=6)
        table.style = "Light Grid Accent 1"

        # Header
        headers = ["Rank", "Shuttle (m³)", "Pump (m³/h)", "NPC (M USD)", "CAPEX (M USD)", "OPEX (M USD)"]
        for col_idx, header in enumerate(headers):
            table.rows[0].cells[col_idx].text = header

        # Data
        for row_idx, (_, row) in enumerate(top10.iterrows(), 1):
            total_capex = (row["NPC_Shuttle_CAPEX_USDm"] +
                          row["NPC_Bunkering_CAPEX_USDm"] +
                          row["NPC_Terminal_CAPEX_USDm"])
            total_opex = (row["NPC_Shuttle_fOPEX_USDm"] + row["NPC_Bunkering_fOPEX_USDm"] +
                         row["NPC_Terminal_fOPEX_USDm"] + row["NPC_Shuttle_vOPEX_USDm"] +
                         row["NPC_Bunkering_vOPEX_USDm"] + row["NPC_Terminal_vOPEX_USDm"])

            table.rows[row_idx].cells[0].text = str(row_idx)
            table.rows[row_idx].cells[1].text = str(int(row["Shuttle_Size_cbm"]))
            table.rows[row_idx].cells[2].text = str(int(row["Pump_Size_m3ph"]))
            table.rows[row_idx].cells[3].text = f"{row['NPC_Total_USDm']:.2f}"
            table.rows[row_idx].cells[4].text = f"{total_capex:.2f}"
            table.rows[row_idx].cells[5].text = f"{total_opex:.2f}"

    def _add_cost_breakdown(self, doc: Document, scenario_df: pd.DataFrame) -> None:
        """Add cost breakdown analysis."""
        doc.add_page_break()
        doc.add_heading("Cost Component Analysis", level=1)

        best = scenario_df.nsmallest(1, "NPC_Total_USDm").iloc[0]

        # Calculate percentages
        total = best["NPC_Total_USDm"]
        capex_pct = ((best["NPC_Shuttle_CAPEX_USDm"] + best["NPC_Bunkering_CAPEX_USDm"] +
                     best["NPC_Terminal_CAPEX_USDm"]) / total * 100)
        opex_pct = 100 - capex_pct

        doc.add_paragraph(
            f"For the optimal solution, the total cost is dominated by "
            f"CAPEX ({capex_pct:.1f}%) and OPEX ({opex_pct:.1f}%)."
        )

        # Breakdown table
        breakdown_table = doc.add_table(rows=5, cols=2)
        breakdown_table.style = "Light Grid Accent 1"

        breakdown = [
            ("Capital Expenditure (CAPEX)", f"{capex_pct:.1f}%"),
            ("Operating Expenditure (OPEX)", f"{opex_pct:.1f}%"),
            ("CAPEX Amount (M USD)", f"{total * capex_pct / 100:.2f}"),
            ("OPEX Amount (M USD)", f"{total * opex_pct / 100:.2f}"),
        ]

        for idx, (item, value) in enumerate(breakdown):
            breakdown_table.rows[idx].cells[0].text = item
            breakdown_table.rows[idx].cells[1].text = str(value)

    def _add_appendix(
        self,
        doc: Document,
        scenario_df: pd.DataFrame,
        yearly_df: pd.DataFrame
    ) -> None:
        """Add appendix with raw data."""
        doc.add_page_break()
        doc.add_heading("Appendix: Raw Data Summary", level=1)

        doc.add_paragraph(f"Total scenarios evaluated: {len(scenario_df)}")
        doc.add_paragraph(f"Time period: {yearly_df['Year'].min():.0f} - {yearly_df['Year'].max():.0f}")

        # Summary statistics
        doc.add_heading("Summary Statistics", level=2)
        stats_table = doc.add_table(rows=6, cols=2)
        stats_table.style = "Light Grid Accent 1"

        stats = [
            ("Minimum NPC (M USD)", f"{scenario_df['NPC_Total_USDm'].min():.2f}"),
            ("Maximum NPC (M USD)", f"{scenario_df['NPC_Total_USDm'].max():.2f}"),
            ("Average NPC (M USD)", f"{scenario_df['NPC_Total_USDm'].mean():.2f}"),
            ("Median NPC (M USD)", f"{scenario_df['NPC_Total_USDm'].median():.2f}"),
            ("Std Dev NPC (M USD)", f"{scenario_df['NPC_Total_USDm'].std():.2f}"),
        ]

        for idx, (stat, value) in enumerate(stats):
            stats_table.rows[idx].cells[0].text = stat
            stats_table.rows[idx].cells[1].text = str(value)

        doc.add_paragraph()
        doc.add_paragraph(
            "For complete data including yearly breakdowns, refer to the accompanying Excel file."
        )
